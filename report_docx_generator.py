# report_docx_generator.py

import os
import re
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def convert_markdown_to_docx(markdown_text, output_path):
    """
    Parses a markdown string and creates a beautifully formatted Word (.docx) document.
    Supports headings, paragraphs, lists, bold text, and tables.
    """
    doc = docx.Document()
    
    # Page setup - 1 inch margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # Split text into lines
    lines = markdown_text.split('\n')
    
    in_table = False
    table_headers = []
    table_rows = []
    
    def process_bold_text(paragraph, text):
        # Parses **bold** and normal text and adds them to paragraph
        parts = re.split(r'(\*\*.*?\*\*)', text)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                run = paragraph.add_run(part[2:-2])
                run.bold = True
            else:
                paragraph.add_run(part)

    def set_cell_background(cell, fill_hex):
        tcPr = cell._tc.get_or_add_tcPr()
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
        tcPr.append(shading)

    def write_table_to_doc():
        if not table_headers:
            return
        num_cols = len(table_headers)
        table = doc.add_table(rows=1, cols=num_cols)
        table.style = 'Light Shading Accent 1'
        
        hdr_cells = table.rows[0].cells
        for i, header in enumerate(table_headers):
            hdr_cells[i].text = ""
            p = hdr_cells[i].paragraphs[0]
            process_bold_text(p, header)
            set_cell_background(hdr_cells[i], "1F497D") # Navy blue header
            for run in p.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255) # white text
        
        for row_data in table_rows:
            row_cells = table.add_row().cells
            for i in range(min(num_cols, len(row_data))):
                row_cells[i].text = ""
                p = row_cells[i].paragraphs[0]
                process_bold_text(p, row_data[i])

    for line in lines:
        stripped = line.strip()
        
        # Detect table start/content
        if stripped.startswith('|'):
            if not in_table:
                in_table = True
                table_headers = [col.strip() for col in stripped.split('|')[1:-1]]
                table_rows = []
            else:
                # Check if it's separator row (e.g. |:---|:---|)
                if re.match(r'^\|[\s:-|]*\|$', stripped):
                    continue
                row_cols = [col.strip() for col in stripped.split('|')[1:-1]]
                table_rows.append(row_cols)
            continue
        else:
            if in_table:
                write_table_to_doc()
                in_table = False
                table_headers = []
                table_rows = []
                
        if not stripped:
            continue
            
        # Detect Headings
        if stripped.startswith('#'):
            h_level = len(stripped) - len(stripped.lstrip('#'))
            title = stripped.lstrip('#').strip()
            
            p = doc.add_heading('', level=min(h_level, 4))
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.keep_with_next = True
            
            run = p.add_run(title)
            run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D) # Navy color for headers
            if h_level == 1:
                run.font.size = Pt(16)
            elif h_level == 2:
                run.font.size = Pt(13)
            else:
                run.font.size = Pt(11)
            continue
            
        # Detect Bullet list items
        if stripped.startswith('- ') or stripped.startswith('* '):
            item_text = stripped[2:].strip()
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_after = Pt(2)
            process_bold_text(p, item_text)
            continue
            
        # Normal paragraphs
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        process_bold_text(p, stripped)
        
    # Flush remaining table if it ends at EOF
    if in_table:
        write_table_to_doc()
        
    # Create target directory if it does not exist
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
